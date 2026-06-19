"""DOCX file reading and writing with formatting preservation."""

import re
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from .config import EAST_ASIAN_FONT, FONT_SIZE_PT, LATIN_FONT


@dataclass
class Run:
    """A text segment with consistent formatting."""
    text: str
    bold: bool = False
    italic: bool = False


@dataclass
class Paragraph:
    """A paragraph containing multiple runs."""
    runs: list[Run]

    @property
    def text(self) -> str:
        """Get the full text of the paragraph."""
        return "".join(run.text for run in self.runs)

    def is_empty(self) -> bool:
        """Check if the paragraph has no text content."""
        return not self.text.strip()


def convert_doc_to_docx(doc_path: Path) -> Path:
    """
    Convert a .doc file to .docx using LibreOffice, returning the output path.

    The output is written to a temporary directory; callers should delete it
    when done.  Raises RuntimeError if LibreOffice is not available or fails.
    """
    tmp_dir = Path(tempfile.mkdtemp())
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(tmp_dir), str(doc_path)],
            capture_output=True,
            text=True,
            timeout=120,
        )
    except FileNotFoundError:
        raise RuntimeError(
            "LibreOffice is required to process .doc files but was not found. "
            "Install it with: sudo apt install libreoffice"
        )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
    converted = tmp_dir / (doc_path.stem + ".docx")
    if not converted.exists():
        raise RuntimeError(f"LibreOffice did not produce expected output: {converted}")
    return converted


def read_docx(file_path: str | Path) -> list[Paragraph]:
    """
    Read a DOCX file and extract paragraphs with formatting.

    Args:
        file_path: Path to the input DOCX file

    Returns:
        List of Paragraph objects with formatting information
    """
    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        runs = []
        for run in para.runs:
            if run.text:  # Only include non-empty runs
                runs.append(Run(
                    text=run.text,
                    bold=run.bold or False,
                    italic=run.italic or False,
                ))

        # Include paragraph even if empty (to preserve structure)
        paragraphs.append(Paragraph(runs=runs))

    return paragraphs


def _set_run_fonts(run) -> None:
    """
    Set fonts for a run using configured LATIN_FONT and EAST_ASIAN_FONT.

    Args:
        run: A python-docx Run object
    """
    # Set Latin/ASCII font
    run.font.name = LATIN_FONT
    run.font.size = Pt(FONT_SIZE_PT)

    # Set East Asian font via XML (python-docx doesn't expose this directly)
    r_element = run._element
    rPr = r_element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), EAST_ASIAN_FONT)


def write_docx(
    paragraphs: list[Paragraph],
    output_path: str | Path,
    translated_texts: list[str],
) -> None:
    """
    Write translated content to a DOCX file, preserving formatting.

    Uses Calibri for English text and Microsoft YaHei for Chinese text.

    Args:
        paragraphs: Original paragraphs with formatting information
        output_path: Path for the output DOCX file
        translated_texts: List of translated text for each paragraph
    """
    doc = Document()

    # Set default font for the Normal style
    style = doc.styles["Normal"]
    style.font.name = LATIN_FONT
    style.font.size = Pt(FONT_SIZE_PT)
    # Set East Asian font for the style
    style_element = style._element
    rPr = style_element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), EAST_ASIAN_FONT)

    for i, (orig_para, translated_text) in enumerate(zip(paragraphs, translated_texts)):
        # Skip empty paragraphs entirely (no empty lines in output)
        if orig_para.is_empty() and not translated_text.strip():
            continue

        para = doc.add_paragraph()

        if translated_text.strip():
            run = para.add_run(translated_text)
            _set_run_fonts(run)

    doc.save(output_path)



def get_plain_text(paragraphs: list[Paragraph]) -> str:
    """
    Get plain text from paragraphs with paragraph markers.

    Args:
        paragraphs: List of Paragraph objects

    Returns:
        Plain text with [P1], [P2], etc. markers
    """
    lines = []
    for i, para in enumerate(paragraphs, 1):
        text = para.text.strip()
        if text:
            lines.append(f"[P{i}] {text}")
        else:
            lines.append(f"[P{i}]")
    return "\n".join(lines)


def parse_translated_text(translated: str, num_paragraphs: int) -> list[str]:
    """
    Parse translated text with paragraph markers back into a list.

    Args:
        translated: Translated text with [P1], [P2], etc. markers
        num_paragraphs: Expected number of paragraphs

    Returns:
        List of translated text for each paragraph
    """
    # Extract text for each paragraph marker
    result = [""] * num_paragraphs

    # Pattern to match [P1], [P2], etc. and capture following text
    pattern = r"\[P(\d+)\]\s*(.*?)(?=\[P\d+\]|$)"
    matches = re.findall(pattern, translated, re.DOTALL)

    for num_str, text in matches:
        idx = int(num_str) - 1
        if 0 <= idx < num_paragraphs:
            result[idx] = text.strip()

    return result
